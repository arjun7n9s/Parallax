"""Convert SOLUTION_APPROACH.md to a formatted .docx for hackathon submission.

Reads the markdown file and produces a Word document with:
- Heading 1/2/3 styles
- Bullet/numbered lists rendered as proper Word lists
- Markdown tables rendered as Word tables
- Code blocks rendered as monospace
- Bold/italic preserved
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_bullet(doc, text, level=0):
    """Add a bullet point with bold/italic parsing."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    _add_inline(p, text)


def add_numbered(doc, text, level=0):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    _add_inline(p, text)


def _add_inline(p, text):
    """Parse **bold**, *italic*, `code` and add as runs to paragraph p."""
    # Tokenize: split on **, *, ` while keeping delimiters
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = p.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        else:
            p.add_run(part)


def add_paragraph_text(doc, text):
    p = doc.add_paragraph()
    _add_inline(p, text)


def add_code_block(doc, code):
    """Add a monospace code block."""
    for line in code.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(9)


def parse_md_table(lines, idx):
    """Parse a markdown table starting at lines[idx]. Return (rows, end_idx)."""
    table_lines = []
    i = idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i])
        i += 1
    if len(table_lines) < 2:
        return None, idx

    # First line is header, second is separator (---|---), rest is body
    header_cells = [c.strip() for c in table_lines[0].strip('|').split('|')]
    # Skip separator line (index 1)
    body_lines = table_lines[2:]

    rows = [header_cells]
    for bl in body_lines:
        cells = [c.strip() for c in bl.strip('|').split('|')]
        rows.append(cells)
    return rows, i


def add_md_table(doc, rows):
    """Add a Word table from parsed markdown rows."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'

    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ''
            cell = table.rows[i].cells[j]
            # Clear default paragraph
            cell.text = ''
            p = cell.paragraphs[0]
            _add_inline(p, cell_text)
            # Make header row bold
            if i == 0:
                for run in p.runs:
                    run.bold = True


def main():
    md_path = Path(r"C:/Users/arjun/Desktop/PSBs/GenAI-Cybersec-hackathon/hackathon-submission/SOLUTION_APPROACH.md")
    docx_path = Path(r"C:/Users/arjun/Desktop/PSBs/GenAI-Cybersec-hackathon/hackathon-submission/SOLUTION_APPROACH.docx")

    md_text = md_path.read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()

    # Set base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Heading
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        # Horizontal rule
        elif stripped == '---':
            # Page break-ish separator
            p = doc.add_paragraph()
            p.add_run('─' * 40).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        # Table
        elif stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1]):
            rows, end_idx = parse_md_table(lines, i)
            if rows:
                add_md_table(doc, rows)
                i = end_idx
                continue
        # Bullet
        elif re.match(r'^\s*[-*]\s+', line):
            indent = (len(line) - len(line.lstrip())) // 2
            text = re.sub(r'^\s*[-*]\s+', '', line)
            add_bullet(doc, text, level=indent)
        # Numbered list
        elif re.match(r'^\s*\d+\.\s+', line):
            indent = (len(line) - len(line.lstrip())) // 2
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            add_numbered(doc, text, level=indent)
        # Empty line
        elif not stripped:
            pass
        # Regular paragraph
        else:
            add_paragraph_text(doc, line)

        i += 1

    doc.save(docx_path)
    print(f"Wrote: {docx_path}")
    print(f"Size: {docx_path.stat().st_size:,} bytes")


if __name__ == '__main__':
    main()
